from flask import Flask, render_template, request, url_for, send_file, redirect
import mysql.connector
import os
from werkzeug.utils import secure_filename
from docx import Document  # 确保安装了 python-docx 库
import PyPDF2  # 确保安装了 PyPDF2 库
from urllib.parse import unquote  # 使用 unquote 进行解码

# 创建 Flask 应用实例
app = Flask(__name__)

# 数据库配置
DB_CONFIG = {
    'host': 'localhost',
    'user': 'pan',
    'password': 'panjun1994',
    'database': 'search_db'
}

# 上传文件夹配置
app.config['UPLOAD_FOLDER'] = '/Users/panjun/search_website/search_website/static/pdfs'  # 确保路径正确
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# 创建上传文件夹（如果不存在）
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def get_db_connection():
    """获取数据库连接"""
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        print("数据库连接成功")  # 添加调试信息
        return conn
    except mysql.connector.Error as err:
        print(f"数据库连接错误: {err}")
        raise

def allowed_file(filename):
    """检查文件是否允许上传"""
    ALLOWED_EXTENSIONS = {'pdf', 'docx'}  # 允许上传的文件类型
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_pdf_content(file_path):
    """从 PDF 文件中提取内容"""
    content = ''
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            content += page.extract_text() + '\n'
    return content.strip()

@app.route('/')
def index():
    """首页路由"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """上传文件"""
    if 'file' not in request.files:
        return '没有文件被上传'
    file = request.files['file']
    if file.filename == '':
        return '没有选择文件'
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # 提取 PDF 内容
        content = extract_pdf_content(file_path) if filename.endswith('.pdf') else '内容来自 Word 文件'

        # 插入数据库的逻辑
        conn = get_db_connection()
        cursor = conn.cursor()
        sql = """
        INSERT INTO articles (title, content, author, created_at, updated_at, effective_date, type, file_name, file_type) 
        VALUES (%s, %s, %s, NOW(), NOW(), %s, %s, %s, %s)
        """

        # 示例数据，您可以根据需要修改
        title = '关于人身损害赔偿项目计算标准的指引'  # 替换为法规标题
        author = '浙江省高级人民法院'  # 替换为作者名
        effective_date = '2022-08-24'  # 使用合适的日期格式
        type = '法规'  # 替换为类型
        file_type = filename.rsplit('.', 1)[1].lower()  # 获取文件类型

        cursor.execute(sql, (title, content, author, effective_date, type, filename, file_type))
        conn.commit()

        return redirect(url_for('search'))  # 上传成功后重定向到搜索结果
    return '文件类型不被支持'

@app.route('/search/')
def search():
    """搜索路由"""
    keyword = request.args.get('keyword', '')
    search_type = request.args.get('type', '')  # 获取搜索类型

    if not keyword:
        return render_template('index.html')

    try:
        # 连接数据库
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # 使用通配符进行模糊搜索
        search_term = f"%{keyword}%"
        sql = """
        SELECT id, title, content, author, effective_date, created_at, type, file_name, file_type 
        FROM articles 
        WHERE (title LIKE %s OR content LIKE %s)
        """

        # 根据搜索类型调整查询
        if search_type:
            sql += " AND type = %s"
            cursor.execute(sql, (search_term, search_term, search_type))
        else:
            cursor.execute(sql, (search_term, search_term))

        results = cursor.fetchall()

        # 处理结果
        for result in results:
            result['created_at'] = result['created_at'].strftime('%Y-%m-%d %H:%M:%S')
            result['effective_date'] = result['effective_date'].strftime('%Y-%m-%d') if result['effective_date'] else None
            result['content'] = result['content'].replace('\r\n', '\n')
            if len(result['content']) > 500:
                result['content'] = result['content'][:500] + '...'

        return render_template('search_results.html',
                               keyword=keyword,
                               results=results)

    except mysql.connector.Error as err:
        print(f"数据库错误: {err}")  # 打印错误信息
        return render_template('search_results.html',
                               keyword=keyword,
                               error=f"数据库错误: {err}")

    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals() and conn.is_connected():
            conn.close()

@app.route('/pdf/<path:filename>')
def serve_pdf(filename):
    """提供PDF文件的路由"""
    # 对文件名进行 URL 解码
    filename = unquote(filename)  # 确保文件名被正确解码
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    print(f"请求的文件路径: {file_path}")  # 打印文件路径
    if not os.path.exists(file_path):
        return f"文件未找到: {file_path}", 404  # 返回404错误
    return send_file(file_path, mimetype='application/pdf', as_attachment=False)

@app.route('/view_pdf/<path:filename>')
def view_pdf(filename):
    """嵌入PDF文件的路由"""
    return render_template('view_pdf.html', filename=filename)

@app.route('/download/<int:article_id>')
def download_article(article_id):
    """下载文章为Word文档"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        sql = "SELECT * FROM articles WHERE id = %s"
        cursor.execute(sql, (article_id,))
        article = cursor.fetchone()

        if article:
            # 创建 Word 文档
            doc = Document()
            doc.add_heading(article['title'], level=1)
            doc.add_paragraph(article['content'])
            # 保存文档到临时文件
            temp_file_path = f"temp_{article_id}.docx"
            doc.save(temp_file_path)
            return send_file(temp_file_path, as_attachment=True, download_name=f"{article['title']}.docx")
        else:
            return "文章不存在", 404

    except Exception as e:
        print(f"下载错误: {e}")
        return f"下载失败: {e}", 500

    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals() and conn.is_connected():
            conn.close()

@app.route('/article/<int:article_id>')
def article_detail(article_id):
    """显示文章详情"""
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    sql = "SELECT * FROM articles WHERE id = %s"
    cursor.execute(sql, (article_id,))
    article = cursor.fetchone()
    if article:
        return render_template('article_detail.html', article=article)
    else:
        return "文章不存在", 404

if __name__ == '__main__':
    app.run(
        host='0.0.0.0',
        port=5001,
        debug=True,
        threaded=True
    )